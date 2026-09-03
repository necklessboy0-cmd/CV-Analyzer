import io
import json
import os
import re
from pathlib import Path
from typing import Any

import streamlit as st
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document


APP_TITLE = "CV Analyzer"
GEMINI_MODEL = "gemini-3.5-flash-lite"
MAX_CV_CHARS = 100_000
SUPPORTED_TYPES = ["pdf", "docx", "txt"]


st.set_page_config(
    page_title=APP_TITLE,
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed",
)


def inject_css() -> None:
    st.markdown(
        """
        <style>
        @keyframes cvFloat {
            0% { transform: translate3d(0, 0, 0) scale(1); }
            50% { transform: translate3d(1.5%, -1%, 0) scale(1.03); }
            100% { transform: translate3d(0, 0, 0) scale(1); }
        }

        .stApp {
            background:
                radial-gradient(circle at 12% 8%, rgba(191, 156, 255, .24), transparent 28%),
                radial-gradient(circle at 88% 18%, rgba(255, 211, 92, .12), transparent 25%),
                linear-gradient(155deg, #f0e7ff 0%, #e9dcff 38%, #24163e 100%);
            color: #120d1c;
        }

        .stApp::before {
            content: "";
            position: fixed;
            inset: -12%;
            pointer-events: none;
            z-index: 0;
            background:
                radial-gradient(circle at 20% 30%, rgba(255,255,255,.18) 0 2px, transparent 3px),
                radial-gradient(circle at 70% 60%, rgba(255,255,255,.13) 0 2px, transparent 3px),
                radial-gradient(circle at 40% 80%, rgba(255,215,105,.10) 0 2px, transparent 3px);
            background-size: 150px 150px, 210px 210px, 180px 180px;
            animation: cvFloat 24s ease-in-out infinite;
        }

        .block-container {
            position: relative;
            z-index: 1;
            max-width: 1180px;
            padding-top: 2.2rem;
            padding-bottom: 3rem;
        }

        .hero {
            padding: 2rem 2.2rem;
            border: 1px solid rgba(255,255,255,.65);
            border-radius: 28px;
            background: rgba(255,255,255,.54);
            box-shadow: 0 22px 70px rgba(32, 15, 63, .20);
            backdrop-filter: blur(14px);
            margin-bottom: 1.5rem;
        }

        .hero h1 {
            margin: 0;
            color: #b38300;
            font-size: clamp(2.3rem, 5vw, 4.2rem);
            letter-spacing: -0.045em;
            font-weight: 900;
        }

        .hero p {
            margin: .55rem 0 0;
            color: #241936;
            font-size: 1.05rem;
        }

        .panel {
            border-radius: 22px;
            padding: 1.2rem 1.35rem;
            background: rgba(24, 13, 43, .76);
            border: 1px solid rgba(255, 220, 122, .18);
            box-shadow: 0 18px 55px rgba(10, 5, 20, .25);
            color: #f7f1ff;
        }

        .panel h3, .panel h4 {
            color: #f3c84f;
        }

        .metric-card {
            border-radius: 20px;
            padding: 1.25rem;
            background: rgba(255,255,255,.92);
            border: 1px solid rgba(62, 34, 95, .12);
            box-shadow: 0 14px 40px rgba(26, 12, 48, .13);
            min-height: 145px;
        }

        .score-number {
            font-size: 3.7rem;
            line-height: 1;
            font-weight: 900;
            color: #5d3297;
        }

        .score-label {
            color: #4c405b;
            font-weight: 700;
        }

        .score-track {
            width: 100%;
            height: 12px;
            border-radius: 99px;
            background: #e6dcf3;
            overflow: hidden;
            margin-top: .85rem;
        }

        .score-fill {
            height: 100%;
            border-radius: 99px;
            background: linear-gradient(90deg, #6d3bb5, #b38300);
        }

        .section-card {
            border-radius: 18px;
            padding: 1.05rem 1.15rem;
            margin: .65rem 0;
            background: rgba(255,255,255,.93);
            border-left: 5px solid #6d3bb5;
            box-shadow: 0 9px 28px rgba(26, 12, 48, .10);
        }

        .section-card strong {
            color: #26143d;
        }

        .issue-card {
            border-radius: 16px;
            padding: .9rem 1rem;
            margin: .55rem 0;
            background: rgba(255,255,255,.93);
            border-left: 4px solid #b38300;
            color: #1c1328;
        }

        .small-note {
            color: #665a73;
            font-size: .88rem;
        }

        div[data-testid="stFileUploader"] {
            border-radius: 18px;
        }

        .stButton > button {
            border-radius: 13px;
            font-weight: 800;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def get_api_key() -> str | None:
    try:
        if "GEMINI_API_KEY" in st.secrets:
            value = st.secrets["GEMINI_API_KEY"]
            if value:
                return str(value).strip()
    except Exception:
        pass

    value = os.getenv("GEMINI_API_KEY")
    return value.strip() if value else None


def read_txt(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("The TXT file could not be decoded as readable text.")


def extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)


def extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    chunks: list[str] = []

    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(text)

    return "\n".join(chunks)


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return extract_pdf(data)
    if suffix == ".docx":
        return extract_docx(data)
    if suffix == ".txt":
        return read_txt(data)

    raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")


def normalize_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text.strip()


def detect_sections(text: str) -> list[str]:
    section_patterns = {
        "Contact Information": r"\b(contact|phone|mobile|email|linkedin|github|portfolio)\b",
        "Professional Summary": r"\b(summary|profile|objective|about me)\b",
        "Experience": r"\b(experience|employment|work history|professional experience)\b",
        "Education": r"\b(education|academic background|qualifications)\b",
        "Skills": r"\b(skills|technical skills|core competencies|competencies)\b",
        "Projects": r"\b(projects|selected projects|academic projects)\b",
        "Certifications": r"\b(certifications|certificates|licenses)\b",
        "Achievements": r"\b(achievements|awards|honors|accomplishments)\b",
        "Languages": r"\b(languages|language proficiency)\b",
    }

    found: list[str] = []
    lower = text.lower()
    for name, pattern in section_patterns.items():
        if re.search(pattern, lower):
            found.append(name)
    return found


def deterministic_checks(text: str) -> dict[str, Any]:
    urls = re.findall(r"(?:https?://|www\.)[^\s<>()]+", text, flags=re.I)
    emails = re.findall(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", text, flags=re.I)

    double_spaces = len(re.findall(r"[ \t]{2,}", text))
    space_before_punctuation = len(re.findall(r"[ \t]+[,.;:!?]", text))
    repeated_punctuation = len(re.findall(r"([!?.,;:])\1+", text))
    trailing_spaces = sum(
        1 for line in text.splitlines() if line.endswith((" ", "\t"))
    )

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    long_lines = sum(1 for line in lines if len(line) > 140)

    return {
        "word_count": len(re.findall(r"\b[\w'-]+\b", text)),
        "character_count": len(text),
        "line_count": len(lines),
        "url_count": len(urls),
        "email_count": len(emails),
        "urls": urls[:50],
        "double_space_count": double_spaces,
        "space_before_punctuation_count": space_before_punctuation,
        "repeated_punctuation_count": repeated_punctuation,
        "trailing_space_line_count": trailing_spaces,
        "very_long_line_count": long_lines,
        "detected_sections": detect_sections(text),
    }


def build_prompt(cv_text: str, checks: dict[str, Any]) -> str:
    return f"""
You are a senior CV/resume auditor and professional career-document editor.

Analyze the supplied CV text deeply and conservatively. The CV is the source of truth.
Do NOT invent education, employment, dates, skills, achievements, links, metrics, employers,
job titles, or personal details. If something is not present, identify it as missing or
unclear rather than assuming it exists.

The application already performed deterministic text checks. Use them as evidence, but do
your own linguistic and professional review as well.

Required audit areas:
1. Overall effectiveness and score from 0 to 100.
2. Executive summary.
3. Strengths.
4. Critical issues.
5. Section-by-section analysis of sections actually present and important missing sections.
6. Micro-level grammar, spelling, punctuation, capitalization, spacing, wording, consistency,
   awkward phrasing, repetition, and unclear statements. Give concrete examples from the CV
   when possible, but keep quotations short.
7. Missing information that would materially improve the CV.
8. ATS/readability/recruiter effectiveness without claiming to have visually rendered the file.
9. Link/contact analysis. Analyze links that are present syntactically; NEVER create or
   fabricate a URL.
10. Actionable recommendations prioritized by impact.
11. Final professional assessment.

Scoring guidance:
- Use the full 0-100 range sensibly.
- Consider clarity, relevance, evidence of impact, experience, skills, education, language
  quality, completeness, consistency, contact information, and recruiter/ATS usability.
- Do not penalize the candidate for sections that are genuinely irrelevant to their profile.
- Do not award points for information that is merely assumed.

Deterministic checks from Python:
{json.dumps(checks, ensure_ascii=False, indent=2)}

Return ONLY valid JSON with exactly this top-level structure:
{{
  "overall_score": 0,
  "score_explanation": "string",
  "executive_summary": "string",
  "strengths": ["string"],
  "critical_issues": ["string"],
  "section_analysis": [
    {{
      "section": "string",
      "status": "strong|needs_improvement|missing|not_applicable",
      "score": 0,
      "findings": ["string"],
      "recommendations": ["string"]
    }}
  ],
  "detailed_errors": [
    {{
      "category": "grammar|spelling|punctuation|spacing|capitalization|wording|consistency|formatting_text|other",
      "severity": "high|medium|low",
      "issue": "string",
      "evidence": "string",
      "fix": "string"
    }}
  ],
  "missing_information": ["string"],
  "recommendations": [
    {{
      "priority": "high|medium|low",
      "recommendation": "string",
      "reason": "string"
    }}
  ],
  "link_analysis": {{
    "present": true,
    "findings": ["string"]
  }},
  "final_assessment": "string"
}}

CV TEXT START
{cv_text}
CV TEXT END
""".strip()


def extract_json(text: str) -> dict[str, Any]:
    cleaned = text.strip()

    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.I)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start >= 0 and end > start:
        parsed = json.loads(cleaned[start : end + 1])
        if isinstance(parsed, dict):
            return parsed

    raise ValueError("Gemini returned an invalid JSON analysis.")


def as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def normalize_result(result: dict[str, Any]) -> dict[str, Any]:
    raw_score = result.get("overall_score", 0)
    try:
        score = int(float(raw_score))
    except (TypeError, ValueError):
        score = 0

    result["overall_score"] = max(0, min(100, score))
    result["score_explanation"] = str(result.get("score_explanation", "")).strip()
    result["executive_summary"] = str(result.get("executive_summary", "")).strip()
    result["strengths"] = [str(x) for x in as_list(result.get("strengths"))]
    result["critical_issues"] = [str(x) for x in as_list(result.get("critical_issues"))]
    result["missing_information"] = [
        str(x) for x in as_list(result.get("missing_information"))
    ]
    result["final_assessment"] = str(result.get("final_assessment", "")).strip()

    sections = []
    for item in as_list(result.get("section_analysis")):
        if not isinstance(item, dict):
            continue
        try:
            item_score = int(float(item.get("score", 0)))
        except (TypeError, ValueError):
            item_score = 0
        sections.append(
            {
                "section": str(item.get("section", "Section")),
                "status": str(item.get("status", "needs_improvement")),
                "score": max(0, min(100, item_score)),
                "findings": [str(x) for x in as_list(item.get("findings"))],
                "recommendations": [
                    str(x) for x in as_list(item.get("recommendations"))
                ],
            }
        )
    result["section_analysis"] = sections

    errors = []
    for item in as_list(result.get("detailed_errors")):
        if not isinstance(item, dict):
            continue
        errors.append(
            {
                "category": str(item.get("category", "other")),
                "severity": str(item.get("severity", "low")),
                "issue": str(item.get("issue", "")),
                "evidence": str(item.get("evidence", "")),
                "fix": str(item.get("fix", "")),
            }
        )
    result["detailed_errors"] = errors

    recommendations = []
    for item in as_list(result.get("recommendations")):
        if isinstance(item, dict):
            recommendations.append(
                {
                    "priority": str(item.get("priority", "medium")),
                    "recommendation": str(item.get("recommendation", "")),
                    "reason": str(item.get("reason", "")),
                }
            )
        else:
            recommendations.append(
                {
                    "priority": "medium",
                    "recommendation": str(item),
                    "reason": "",
                }
            )
    result["recommendations"] = recommendations

    link_analysis = result.get("link_analysis")
    if not isinstance(link_analysis, dict):
        link_analysis = {"present": False, "findings": []}
    link_analysis["present"] = bool(link_analysis.get("present", False))
    link_analysis["findings"] = [
        str(x) for x in as_list(link_analysis.get("findings"))
    ]
    result["link_analysis"] = link_analysis

    return result


def analyze_with_gemini(cv_text: str, checks: dict[str, Any], api_key: str) -> dict[str, Any]:
    client = genai.Client(api_key=api_key)
    prompt = build_prompt(cv_text, checks)

    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
        ),
    )

    response_text = getattr(response, "text", None)
    if not response_text:
        raise ValueError("Gemini returned an empty analysis.")

    return normalize_result(extract_json(response_text))


def show_list(title: str, items: list[str], empty_text: str = "None identified.") -> None:
    st.markdown(f"### {title}")
    if not items:
        st.info(empty_text)
        return
    for item in items:
        st.markdown(f'<div class="issue-card">{item}</div>', unsafe_allow_html=True)


def render_analysis(result: dict[str, Any], checks: dict[str, Any]) -> None:
    score = result["overall_score"]

    st.markdown("## Analysis Result")

    col1, col2, col3 = st.columns([1.05, 1.35, 1.35])
    with col1:
        st.markdown(
            f"""
            <div class="metric-card">
                <div class="score-number">{score}<span style="font-size:1.5rem;">/100</span></div>
                <div class="score-label">Overall CV Score</div>
                <div class="score-track"><div class="score-fill" style="width:{score}%"></div></div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with col2:
        st.markdown(
            f"""
            <div class="metric-card">
                <div style="font-size:.85rem;font-weight:800;color:#6d3bb5;">DOCUMENT</div>
                <div style="font-size:1.7rem;font-weight:900;color:#20142d;">{checks["word_count"]:,} words</div>
                <div class="small-note">{checks["character_count"]:,} characters · {checks["line_count"]:,} text lines</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with col3:
        st.markdown(
            f"""
            <div class="metric-card">
                <div style="font-size:.85rem;font-weight:800;color:#b38300;">CONTACT / LINKS</div>
                <div style="font-size:1.7rem;font-weight:900;color:#20142d;">{checks["email_count"]} email · {checks["url_count"]} link(s)</div>
                <div class="small-note">Detected sections: {len(checks["detected_sections"])}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("---")

    st.markdown("### Score Explanation")
    st.markdown(f'<div class="section-card">{result["score_explanation"]}</div>', unsafe_allow_html=True)

    st.markdown("### Executive Summary")
    st.markdown(f'<div class="section-card">{result["executive_summary"]}</div>', unsafe_allow_html=True)

    show_list("Strengths", result["strengths"], "No major strengths were returned.")
    show_list("Critical Issues", result["critical_issues"], "No critical issues were identified.")

    st.markdown("### Section-by-Section Analysis")
    if result["section_analysis"]:
        for section in result["section_analysis"]:
            status = section["status"].replace("_", " ").title()
            findings = "".join(f"<li>{x}</li>" for x in section["findings"])
            recs = "".join(f"<li>{x}</li>" for x in section["recommendations"])
            st.markdown(
                f"""
                <div class="section-card">
                    <div style="font-size:1.15rem;font-weight:900;">{section["section"]}</div>
                    <div style="margin:.3rem 0 .7rem;"><strong>Status:</strong> {status}
                    &nbsp; | &nbsp; <strong>Score:</strong> {section["score"]}/100</div>
                    {"<strong>Findings</strong><ul>" + findings + "</ul>" if findings else ""}
                    {"<strong>Recommendations</strong><ul>" + recs + "</ul>" if recs else ""}
                </div>
                """,
                unsafe_allow_html=True,
            )
    else:
        st.info("No section analysis was returned.")

    st.markdown("### Micro-Level Errors")
    if result["detailed_errors"]:
        for error in result["detailed_errors"]:
            st.markdown(
                f"""
                <div class="issue-card">
                    <strong>{error["category"].title()} · {error["severity"].title()}</strong><br>
                    <b>Issue:</b> {error["issue"]}<br>
                    <b>Evidence:</b> {error["evidence"]}<br>
                    <b>Fix:</b> {error["fix"]}
                </div>
                """,
                unsafe_allow_html=True,
            )
    else:
        st.success("No specific micro-level errors were returned by the AI audit.")

    show_list(
        "Missing Information",
        result["missing_information"],
        "No material missing information was identified.",
    )

    st.markdown("### Recommendations")
    if result["recommendations"]:
        for item in result["recommendations"]:
            st.markdown(
                f"""
                <div class="section-card">
                    <strong>{item["priority"].title()} Priority</strong><br>
                    {item["recommendation"]}
                    {"<br><span class='small-note'>" + item["reason"] + "</span>" if item["reason"] else ""}
                </div>
                """,
                unsafe_allow_html=True,
            )
    else:
        st.info("No recommendations were returned.")

    st.markdown("### Link Analysis")
    findings = result["link_analysis"]["findings"]
    if findings:
        for item in findings:
            st.markdown(f'<div class="issue-card">{item}</div>', unsafe_allow_html=True)
    else:
        st.info("No link-specific findings were returned.")

    st.markdown("### Automated Text Diagnostics")
    diag = [
        ("Double spaces", checks["double_space_count"]),
        ("Space before punctuation", checks["space_before_punctuation_count"]),
        ("Repeated punctuation", checks["repeated_punctuation_count"]),
        ("Trailing-space lines", checks["trailing_space_line_count"]),
        ("Very long lines", checks["very_long_line_count"]),
    ]
    dcols = st.columns(5)
    for col, (label, value) in zip(dcols, diag):
        with col:
            st.metric(label, value)

    st.markdown("### Final Assessment")
    st.markdown(
        f'<div class="panel"><h3>Professional Conclusion</h3><p>{result["final_assessment"]}</p></div>',
        unsafe_allow_html=True,
    )


def main() -> None:
    inject_css()

    st.markdown(
        """
        <div class="hero">
            <h1>CV Analyzer</h1>
            <p>Deep CV auditing with micro-level language checks, section analysis, scoring, and actionable recommendations.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="panel">
            <h3>Upload your CV</h3>
            <p>Supported formats: PDF, DOCX, TXT. The document is extracted locally in the app,
            then analyzed with Gemini. No external API is used other than Gemini, and the app
            does not generate external links.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    uploaded = st.file_uploader(
        "Choose a CV file",
        type=SUPPORTED_TYPES,
        help="Upload a readable PDF, DOCX, or TXT CV.",
    )

    if not uploaded:
        st.info("Upload a CV to begin the analysis.")
        return

    data = uploaded.getvalue()

    if not data:
        st.error("The uploaded file is empty.")
        return

    if len(data) > 12 * 1024 * 1024:
        st.error("The file is too large. Please upload a CV smaller than 12 MB.")
        return

    if st.button("Analyze CV", type="primary", use_container_width=True):
        api_key = get_api_key()
        if not api_key:
            st.error(
                "Gemini API key not found. Add GEMINI_API_KEY to Streamlit Secrets "
                "or the GEMINI_API_KEY environment variable."
            )
            return

        try:
            with st.spinner("Extracting and auditing the CV..."):
                text = normalize_text(extract_text(uploaded.name, data))

            if not text:
                st.error(
                    "No readable text was extracted. If this is a scanned/image-only PDF, "
                    "use a text-based PDF or DOCX/TXT version."
                )
                return

            if len(text) > MAX_CV_CHARS:
                text = text[:MAX_CV_CHARS]
                st.warning(
                    f"The extracted CV text exceeded {MAX_CV_CHARS:,} characters, "
                    "so the analysis used the first portion of the document."
                )

            checks = deterministic_checks(text)

            with st.spinner("Gemini is performing the deep CV audit..."):
                result = analyze_with_gemini(text, checks, api_key)

            st.session_state["cv_result"] = result
            st.session_state["cv_checks"] = checks
            st.session_state["cv_filename"] = uploaded.name

        except Exception as exc:
            message = str(exc).lower()

            if "429" in message or "rate" in message or "quota" in message:
                st.error("Gemini rate limit or quota reached. Please try again later.")
            elif "401" in message or "403" in message or "api key" in message:
                st.error("Gemini authentication failed. Check GEMINI_API_KEY.")
            elif "json" in message:
                st.error("Gemini returned an unreadable analysis format. Please try again.")
            elif "pdf" in message or "document" in message or "decode" in message:
                st.error("The document could not be read. Try exporting it again as PDF, DOCX, or TXT.")
            else:
                st.error("The CV could not be analyzed. Please verify the file and try again.")

    if "cv_result" in st.session_state and "cv_checks" in st.session_state:
        filename = st.session_state.get("cv_filename", "uploaded CV")
        st.caption(f"Showing analysis for: {filename}")
        render_analysis(st.session_state["cv_result"], st.session_state["cv_checks"])


if __name__ == "__main__":
    main()
